from langchain_core.tools import Tool, tool
import requests
import json
from docx import Document
import re
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage



@tool
def search_pypi(package_name: str) -> str:
    """Search PyPI for Python package information. Input should be the package name.
    Args: 
        package_name: name of the package 
    """
    print(f"Tool called for package: {package_name}")
    base_url = "https://pypi.org/pypi"
    try:
        try:
            response = requests.get(f"{base_url}/{package_name}/json")
            response.raise_for_status()
            info = response.json()
        except requests.RequestException as e:
            raise Exception(f"Error fetching PyPI info for {package_name}: {str(e)}")
        result =  json.dumps({
            "name": info["info"]["name"],
            "summary": info["info"]["summary"],
        })
        print(f"Tool result: {result}")
        return result
    except Exception as e:
        return f"Could not find package information: {str(e)}"
    
# @tool
# def write_to_docx(documentation_text: str) -> str:
#     """
#     Writes the AI-generated documentation to a .docx file and returns the file path.
#     """
#     doc = Document()
#     # doc.add_heading("Code Documentation", level=1)

#     lines = documentation_text.split("\n")
#     for line in lines:
#         if line.startswith("# "):  # Section Heading
#             doc.add_heading(line[2:], level=1)
#         elif line.startswith("## "):  # Subsection Heading
#             doc.add_heading(line[3:], level=2)
#         else:  # Normal paragraph
#             doc.add_paragraph(line)

#     file_path = "generated_documentation.docx"
#     doc.save(file_path)
#     return file_path

def write_to_docx(state):
    text = state['messages'][-1].content
    filename = 'generated_documentation.docx'
    doc = Document()
    
    lines = text.split("\n")
    for line in lines:
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif "**" in line:
            bold_parts = re.split(r"(\*\*.*?\*\*)", line)
            para = doc.add_paragraph()
            for part in bold_parts:
                if part.startswith("**") and part.endswith("**"):
                    para.add_run(part[2:-2]).bold = True
                else:
                    para.add_run(part)
        else:
            doc.add_paragraph(line)
    
    # Save document
    doc.save(filename)
    return {"messages": [SystemMessage(content=[filename])]}